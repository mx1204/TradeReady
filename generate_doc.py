from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x37, 0x6C)   # headings
TEAL   = RGBColor(0x00, 0x87, 0x8A)   # accent / sub-headings
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF7, 0xF8)   # table row shading
DARK   = RGBColor(0x1A, 0x37, 0x6C)   # table header bg

# hex strings for shd element (RGBColor doesn't expose .red/.green/.blue)
NAVY_HEX  = "1A376C"
TEAL_HEX  = "00878A"
WHITE_HEX = "FFFFFF"
LIGHT_HEX = "F0F7F8"
DARK_HEX  = "1A376C"

# ── Helper functions ──────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(text, level=1, colour=NAVY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    run.font.color.rgb = colour
    return p

def add_body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "6")
    bottom.set(qn("w:space"),"1")
    bottom.set(qn("w:color"),"008788")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
# COVER / TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TradeReady")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-Powered Cross-Border Trade Compliance")
r.font.size = Pt(16)
r.font.color.rgb = TEAL
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Singapore  ↔  Malaysia  |  Electronics Exports & Imports")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
add_divider()
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. The Problem", level=1)
add_body(
    "Cross-border electronics trade between Singapore and Malaysia is growing rapidly, "
    "yet the compliance process remains slow, manual, and error-prone. Sellers must "
    "navigate two separate customs regimes — each with its own HS tariff codes, duty rates, "
    "tax rules, certification requirements, and documentation checklists."
)
add_body("This creates real pain for businesses of all sizes:")

bullets_problem = [
    ("Complex regulations: ", "HS codes, duty/tax rates, and certification requirements differ between Malaysia and Singapore."),
    ("High stakes: ", "Incorrect declarations lead to delays, fines, or shipment seizure."),
    ("Manual research: ", "Sellers must read official customs portals and interpret legal language themselves."),
    ("Expensive consultants: ", "Hiring compliance experts costs thousands of dollars and adds days of turnaround time."),
    ("No verification: ", "There is no easy way to double-check a declaration before it is submitted."),
]
for bp, bt in bullets_problem:
    add_bullet(bt, bold_prefix=bp)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. OUR SOLUTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Our Solution — TradeReady", level=1)
add_body(
    "TradeReady is an AI-powered compliance assistant that turns a product photo and "
    "basic shipment details into a verified, ready-to-submit customs declaration — in seconds."
)
add_body(
    "It combines OpenAI vision, a multi-agent AI pipeline, and a curated evidence store "
    "of official customs rules to automatically classify products, calculate duties and taxes, "
    "identify required certifications, and pre-fill declaration forms — all with every claim "
    "traceable back to an official source."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. WHO IS IT FOR
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Who Is It For", level=1)

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.columns[0].width = Cm(7)
tbl.columns[1].width = Cm(10)

hdr = tbl.rows[0].cells
set_cell_bg(hdr[0], DARK_HEX); set_cell_bg(hdr[1], DARK_HEX)
for i, txt in enumerate(["User", "Description"]):
    r = hdr[i].paragraphs[0].add_run(txt)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(11)

rows_who = [
    ("E-commerce sellers", "SMEs selling electronics on Shopee, Lazada, or direct channels shipping between SG and MY."),
    ("Logistics providers", "Freight forwarders and 3PLs that prepare customs paperwork on behalf of their clients."),
    ("Customs brokers", "Compliance professionals who need faster, auditable classification and duty estimates."),
]
for i, (u, d) in enumerate(rows_who):
    row = tbl.add_row().cells
    if i % 2 == 0:
        set_cell_bg(row[0], LIGHT_HEX); set_cell_bg(row[1], LIGHT_HEX)
    for j, txt in enumerate([u, d]):
        r = row[j].paragraphs[0].add_run(txt)
        r.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. HOW IT WORKS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. How It Works", level=1)
add_body("TradeReady follows a six-step workflow:")

steps = [
    ("Step 1 — Product Identification",
     "The seller uploads a photo of their product. TradeReady uses OpenAI Vision to automatically detect the product category (e.g. wireless earbuds, smartwatch) with a confidence score. The seller confirms the result before any compliance work begins."),
    ("Step 2 — Shipment Details",
     "The seller provides basic shipment information: quantity, unit value, currency, origin, destination, and invoice number. These can be typed or captured by voice on the frontend."),
    ("Step 3 — Multi-Agent Compliance Generation",
     "Five specialist AI agents run in sequence to build the full compliance package (see Section 5)."),
    ("Step 4 — Critic Verification",
     "An independent Critic Agent cross-checks every generated output against the evidence store — recomputing duty math, verifying HS code citations, confirming certification requirements. Only verified outputs are returned as ready-to-submit."),
    ("Step 5 — Auto-Filled Declaration",
     "TradeReady delivers a pre-filled declaration form with all verified fields: HS code, duty/tax amounts, required documents, restriction status, and certification bodies."),
    ("Step 6 — Human Escalation (When Needed)",
     "If the Critic flags an issue — such as a missing evidence source or a maths mismatch — the case is escalated to a human compliance reviewer. Most transactions pass automatically."),
]
for title, body in steps:
    add_bullet(body, bold_prefix=title + ": ")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. THE AI AGENT PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. The AI Agent Pipeline", level=1)
add_body(
    "The core of TradeReady is a multi-agent architecture where each agent is a specialist. "
    "This prevents any single model from hallucinating across too many domains and makes "
    "every output independently verifiable."
)

agents_tbl = doc.add_table(rows=1, cols=3)
agents_tbl.style = "Table Grid"
agents_tbl.columns[0].width = Cm(4.5)
agents_tbl.columns[1].width = Cm(5.5)
agents_tbl.columns[2].width = Cm(7)

hdr = agents_tbl.rows[0].cells
set_cell_bg(hdr[0], DARK_HEX); set_cell_bg(hdr[1], DARK_HEX); set_cell_bg(hdr[2], DARK_HEX)
for i, txt in enumerate(["Agent", "Role", "Output"]):
    r = hdr[i].paragraphs[0].add_run(txt)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(11)

agents_data = [
    ("Classification Agent", "Maps product to HS & local tariff codes", "HS 6-digit code + jurisdiction code + confidence"),
    ("Duty/Tax Agent",       "Calculates customs value, import duty, tax", "Exact duty & tax amounts + landed cost estimate"),
    ("Restriction Agent",    "Checks certifications & controlled goods rules", "Restriction status + required certification bodies (SIRIM/IMDA)"),
    ("Document Agent",       "Generates docs checklist & auto-fills form fields", "Required documents list + pre-filled declaration fields"),
    ("Critic Agent",         "Independent cross-check of all outputs vs evidence", "Pass / Human-review-required verdict + issues list"),
]
for i, (a, r_txt, o) in enumerate(agents_data):
    row = agents_tbl.add_row().cells
    if i % 2 == 0:
        set_cell_bg(row[0], LIGHT_HEX); set_cell_bg(row[1], LIGHT_HEX); set_cell_bg(row[2], LIGHT_HEX)
    for j, txt in enumerate([a, r_txt, o]):
        run = row[j].paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        if j == 0:
            run.bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. KEY OUTPUTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6. What TradeReady Delivers", level=1)
add_body("A single compliance run produces:")

outputs = [
    ("HS code classification", " with rationale and official source citation"),
    ("Duty & tax breakdown", " (customs value → import duty → tax → landed cost)"),
    ("Restriction status", " and which certification bodies are required (e.g. SIRIM, IMDA)"),
    ("Required documents checklist", " tailored to the destination jurisdiction"),
    ("Auto-filled declaration form", " ready to paste into Shopee, Lazada, or a customs portal"),
    ("Evidence pack", " — every claim linked to an official customs rule"),
    ("Critic report", " — a pass/fail audit of all generated outputs"),
    ("Jurisdiction comparison", " — instant diff if the seller switches from Malaysia to Singapore or vice versa"),
]
for bold, rest in outputs:
    add_bullet(rest, bold_prefix=bold + ":")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. WHY TRADEREADY IS DIFFERENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("7. Why TradeReady Is Different", level=1)

diff_tbl = doc.add_table(rows=1, cols=2)
diff_tbl.style = "Table Grid"
diff_tbl.columns[0].width = Cm(8)
diff_tbl.columns[1].width = Cm(9)

hdr = diff_tbl.rows[0].cells
set_cell_bg(hdr[0], DARK_HEX); set_cell_bg(hdr[1], DARK_HEX)
for i, txt in enumerate(["Differentiator", "What It Means in Practice"]):
    r = hdr[i].paragraphs[0].add_run(txt)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(11)

diff_data = [
    ("Source-grounded outputs",
     "Every claim is backed by an official customs rule. Users see exactly where each requirement comes from — no black-box answers."),
    ("Critic agent cross-check",
     "An independent AI agent re-verifies every output before it is returned. If maths is wrong or a citation is missing, it is caught before the seller sees it."),
    ("Human-on-the-loop (not in-the-loop)",
     "Humans are a safety net, not a bottleneck. High-confidence cases auto-complete; only genuine edge cases escalate."),
    ("Jurisdiction-aware",
     "Malaysia and Singapore have different HS codes, tax rates, and certification regimes. TradeReady handles both natively, not as a one-size-fits-all answer."),
    ("Photo → compliance in one flow",
     "No manual product lookup. The seller uploads a photo; the AI identifies the product and runs the full compliance pipeline end to end."),
]
for i, (d, m) in enumerate(diff_data):
    row = diff_tbl.add_row().cells
    if i % 2 == 0:
        set_cell_bg(row[0], LIGHT_HEX); set_cell_bg(row[1], LIGHT_HEX)
    r0 = row[0].paragraphs[0].add_run(d); r0.bold = True; r0.font.size = Pt(11)
    r1 = row[1].paragraphs[0].add_run(m); r1.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 8. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
add_heading("8. Technology Stack", level=1)

tech_bullets = [
    ("Backend: ", "FastAPI (Python) — lightweight, async-ready REST API"),
    ("AI Model: ", "OpenAI GPT-4.1-mini Vision — product identification from photos"),
    ("Agent Framework: ", "Custom multi-agent orchestrator with five specialist agents"),
    ("Evidence Store: ", "Curated RAG dataset of Malaysia & Singapore electronics customs rules"),
    ("Frontend: ", "Responsive HTML/JS UI with voice input support"),
    ("Deployment: ", "Single PowerShell launcher starts both backend (port 8000) and frontend (port 4174)"),
]
for bp, bt in tech_bullets:
    add_bullet(bt, bold_prefix=bp)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 9. CURRENT SCOPE & NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("9. Current Scope & Roadmap", level=1)
add_body("MVP covers four high-volume electronics categories:")
for cat in ["Wireless earbuds", "Bluetooth speakers", "Smartwatches", "Phone chargers"]:
    add_bullet(cat)

doc.add_paragraph()
add_body("Planned extensions:", bold=True)
roadmap = [
    "Expand evidence store to all electronics HS chapters (84, 85, 90)",
    "Add Thailand, Indonesia, and Vietnam as destination jurisdictions",
    "Live sync with official customs tariff portals via API",
    "Seller dashboard with run history, comparison reports, and export to PDF",
    "Platform integrations (Shopee, Lazada, EasyParcel) for one-click form submission",
]
for r in roadmap:
    add_bullet(r)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 10. KNOWN LIMITATIONS (HONESTY SECTION)
# ══════════════════════════════════════════════════════════════════════════════
add_heading("10. Known Limitations", level=1)
add_body(
    "TradeReady is a proof-of-concept built for this hackathon. We want to be transparent "
    "about what it is and is not yet:"
)
limits = [
    "Evidence store is a demo dataset — rates and codes must be verified against live official portals before real submission.",
    "MVP product scope is limited to four categories; the architecture supports any product.",
    "Photos confirm product type only — quantity, value, and origin must still be supplied by the seller.",
    "No live integration with customs authority systems (MySST, TradeNet) in this version.",
]
for l in limits:
    add_bullet(l)

doc.add_paragraph()
add_divider()
doc.add_paragraph()

# ── Footer tagline ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TradeReady — Compliance confidence, from photo to declaration.")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = TEAL

# ── Save ───────────────────────────────────────────────────────────────────────
out = "TradeReady_Product_Pitch.docx"
doc.save(out)
print(f"Saved: {out}")
